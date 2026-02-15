"""Word 문서 생성 서비스"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from typing import List, Dict, Any


def generate_safe_contract(
    contract_type: str,
    clauses: List[Dict[str, Any]],
    apply_alternatives: bool = True
) -> BytesIO:
    """
    분석된 조항들을 기반으로 수정된 계약서 Word 문서 생성

    Args:
        contract_type: 계약서 유형 (예: "투자계약서")
        clauses: 분석된 조항 리스트
        apply_alternatives: True면 alternative 적용, False면 원본 유지

    Returns:
        BytesIO: Word 문서 바이트 스트림
    """
    doc = Document()

    # 제목 추가
    title = doc.add_heading(contract_type, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 부제: 수정 버전 표시
    if apply_alternatives:
        subtitle = doc.add_paragraph("(AI 검토 수정본)")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in subtitle.runs:
            run.italic = True
            run.font.size = Pt(11)

    doc.add_paragraph()  # 빈 줄

    modified_count = 0

    # 각 조항 추가
    for clause in clauses:
        # 조항 제목
        clause_title = clause.get("title", f"제{clause.get('number', '')}조")
        doc.add_heading(clause_title, level=1)

        # 조항 내용 결정
        analysis = clause.get("analysis", {})
        risk_score = analysis.get("risk_score", 0) if isinstance(analysis, dict) else 0
        alternative = clause.get("alternative", "")
        original = clause.get("content", "")

        if apply_alternatives and alternative and risk_score >= 7:
            # 수정된 내용 사용
            content = alternative
            modified_count += 1

            # 수정 표시
            note = doc.add_paragraph()
            note_run = note.add_run("[AI 수정 적용]")
            note_run.italic = True
            note_run.font.size = Pt(9)
        else:
            content = original

        # 내용 추가
        para = doc.add_paragraph(content)
        para.paragraph_format.first_line_indent = Inches(0.3)

        doc.add_paragraph()  # 조항 간 간격

    # 문서 끝에 수정 요약 추가
    if apply_alternatives and modified_count > 0:
        doc.add_paragraph()
        doc.add_paragraph("─" * 40)
        summary = doc.add_paragraph()
        summary_run = summary.add_run(
            f"본 문서는 AI가 {modified_count}개 조항을 검토하여 수정한 버전입니다.\n"
            "수정된 조항은 [AI 수정 적용] 표시가 되어 있습니다.\n"
            "최종 계약 체결 전 법률 전문가의 검토를 권장합니다."
        )
        summary_run.italic = True
        summary_run.font.size = Pt(9)

    # 문서를 BytesIO로 저장
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
